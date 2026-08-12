"""Format-specific text extraction for the AI Policy Builder
(AI_EXTRACTION_PIPELINE.md Stage 2). Every supported format is normalized
to one plain-text blob with inline location markers, so the LLM analysis
stage (claude_provider.py) never needs to know what format the source
document was.

This module has no knowledge of RuntimePolicy, the extraction provider,
or the database; it only turns bytes into marked-up text, the same
narrow-responsibility discipline domain/runtime_policy/ already holds
itself to.
"""

import csv
import io
from dataclasses import dataclass

import openpyxl
from docx import Document as DocxDocument
from pypdf import PdfReader

SUPPORTED_FORMATS = ("pdf", "docx", "xlsx", "csv", "text")


@dataclass(frozen=True)
class CoverageStats:
    """Coverage Analysis (Authority Intelligence Program, Phase 3,
    EXPLAINABILITY_MODEL.md): deterministic counts produced by the
    parser itself, never an LLM's self-report of how much of a document
    it "covered." A `clause` here means whatever this format's own
    location marker in extract_text() denotes -- a PDF page, a docx
    paragraph/table row, an xlsx/csv row."""

    clauses_analysed: int
    clauses_ignored: int
    tables_extracted: int
    images_skipped: int

    @property
    def coverage_percent(self) -> float:
        total = self.clauses_analysed + self.clauses_ignored
        return round(100.0 * self.clauses_analysed / total, 1) if total else 100.0

_EXTENSION_TO_FORMAT = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".xlsx": "xlsx",
    ".xls": "xlsx",
    ".csv": "csv",
    ".txt": "text",
}

_CONTENT_TYPE_TO_FORMAT = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet": "xlsx",
    "application/vnd.ms-excel": "xlsx",
    "text/csv": "csv",
    "text/plain": "text",
}


class UnsupportedFormatError(Exception):
    pass


def detect_format(filename: str, content_type: str | None) -> str:
    """Filename extension first (more reliable than a browser-supplied
    content_type, which is frequently generic or wrong for these formats),
    falling back to content_type."""
    lower_name = (filename or "").lower()
    for ext, fmt in _EXTENSION_TO_FORMAT.items():
        if lower_name.endswith(ext):
            return fmt
    if content_type in _CONTENT_TYPE_TO_FORMAT:
        return _CONTENT_TYPE_TO_FORMAT[content_type]
    raise UnsupportedFormatError(f"unsupported_format: {filename!r} ({content_type!r})")


def _extract_pdf(raw: bytes) -> str:
    reader = PdfReader(io.BytesIO(raw))
    parts = []
    for i, page in enumerate(reader.pages):
        parts.append(f"--- page {i + 1} ---\n{page.extract_text() or ''}")
    return "\n\n".join(parts)


def _extract_docx(raw: bytes) -> str:
    doc = DocxDocument(io.BytesIO(raw))
    parts = []
    n = 0
    for para in doc.paragraphs:
        if not para.text.strip():
            continue
        n += 1
        parts.append(f"--- paragraph {n} ---\n{para.text}")
    for table in doc.tables:
        for row in table.rows:
            n += 1
            cells = [c.text for c in row.cells]
            parts.append(f"--- paragraph {n} ---\n{chr(9).join(cells)}")
    return "\n\n".join(parts)


def _extract_xlsx(raw: bytes) -> str:
    workbook = openpyxl.load_workbook(io.BytesIO(raw), data_only=True, read_only=True)
    parts = []
    for sheet in workbook.worksheets:
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if all(cell is None for cell in row):
                continue
            cells = ["" if cell is None else str(cell) for cell in row]
            parts.append(f"--- sheet '{sheet.title}', row {row_idx} ---\n{chr(9).join(cells)}")
    return "\n\n".join(parts)


def _extract_csv(raw: bytes) -> str:
    text = raw.decode("utf-8", errors="replace")
    parts = []
    for row_idx, row in enumerate(csv.reader(io.StringIO(text)), start=1):
        if not any(cell.strip() for cell in row):
            continue
        parts.append(f"--- row {row_idx} ---\n{chr(9).join(row)}")
    return "\n\n".join(parts)


def _extract_text(raw: bytes) -> str:
    text = raw.decode("utf-8", errors="replace")
    return f"--- document ---\n{text}"


_EXTRACTORS = {
    "pdf": _extract_pdf,
    "docx": _extract_docx,
    "xlsx": _extract_xlsx,
    "csv": _extract_csv,
    "text": _extract_text,
}


def extract_text(format: str, raw: bytes) -> str:
    """Returns marked-up plain text. An empty or all-blank document
    produces an empty string, a valid outcome (AI_EXTRACTION_PIPELINE.md
    Stage 2), not an error; callers treat zero candidates from empty text
    as a normal, successfully extracted (zero-result) upload."""
    if format not in _EXTRACTORS:
        raise UnsupportedFormatError(f"unsupported_format: {format!r}")
    return _EXTRACTORS[format](raw)


def _extract_pdf_with_coverage(raw: bytes) -> tuple[str, CoverageStats]:
    reader = PdfReader(io.BytesIO(raw))
    parts = []
    analysed = ignored = images = 0
    for i, page in enumerate(reader.pages):
        text = page.extract_text() or ""
        parts.append(f"--- page {i + 1} ---\n{text}")
        if text.strip():
            analysed += 1
        else:
            # A page with no extractable text -- most commonly a scanned
            # image with no text layer. Genuinely unsupported by this
            # extractor (no OCR here), not silently treated as "empty and
            # fine": it's counted so a reviewer can see it happened.
            ignored += 1
        try:
            images += len(page.images)
        except Exception:
            pass
    return "\n\n".join(parts), CoverageStats(analysed, ignored, 0, images)


def _extract_docx_with_coverage(raw: bytes) -> tuple[str, CoverageStats]:
    doc = DocxDocument(io.BytesIO(raw))
    parts = []
    n = analysed = ignored = 0
    for para in doc.paragraphs:
        if not para.text.strip():
            ignored += 1
            continue
        n += 1
        analysed += 1
        parts.append(f"--- paragraph {n} ---\n{para.text}")
    for table in doc.tables:
        for row in table.rows:
            n += 1
            analysed += 1
            cells = [c.text for c in row.cells]
            parts.append(f"--- paragraph {n} ---\n{chr(9).join(cells)}")
    images_skipped = len(doc.inline_shapes)
    return "\n\n".join(parts), CoverageStats(analysed, ignored, len(doc.tables), images_skipped)


def _extract_xlsx_with_coverage(raw: bytes) -> tuple[str, CoverageStats]:
    workbook = openpyxl.load_workbook(io.BytesIO(raw), data_only=True, read_only=True)
    parts = []
    analysed = ignored = 0
    for sheet in workbook.worksheets:
        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            if all(cell is None for cell in row):
                ignored += 1
                continue
            analysed += 1
            cells = ["" if cell is None else str(cell) for cell in row]
            parts.append(f"--- sheet '{sheet.title}', row {row_idx} ---\n{chr(9).join(cells)}")
    return "\n\n".join(parts), CoverageStats(analysed, ignored, len(workbook.worksheets), 0)


def _extract_csv_with_coverage(raw: bytes) -> tuple[str, CoverageStats]:
    text = raw.decode("utf-8", errors="replace")
    parts = []
    analysed = ignored = 0
    for row_idx, row in enumerate(csv.reader(io.StringIO(text)), start=1):
        if not any(cell.strip() for cell in row):
            ignored += 1
            continue
        analysed += 1
        parts.append(f"--- row {row_idx} ---\n{chr(9).join(row)}")
    return "\n\n".join(parts), CoverageStats(analysed, ignored, 1, 0)


def _extract_text_with_coverage(raw: bytes) -> tuple[str, CoverageStats]:
    text = raw.decode("utf-8", errors="replace")
    analysed = 1 if text.strip() else 0
    return f"--- document ---\n{text}", CoverageStats(analysed, 1 - analysed, 0, 0)


_EXTRACTORS_WITH_COVERAGE = {
    "pdf": _extract_pdf_with_coverage,
    "docx": _extract_docx_with_coverage,
    "xlsx": _extract_xlsx_with_coverage,
    "csv": _extract_csv_with_coverage,
    "text": _extract_text_with_coverage,
}


def extract_text_with_coverage(format: str, raw: bytes) -> tuple[str, CoverageStats]:
    """Coverage Analysis (Phase 3): same marked-up text as extract_text(),
    plus deterministic parsing statistics. A NEW function, not a change
    to extract_text()'s signature -- the original AI Policy Builder's
    single-document upload path keeps calling extract_text() completely
    unchanged; only the AI Authority Builder's corpus path (Phase 3) calls
    this one, so there is no behavior change for any existing caller."""
    if format not in _EXTRACTORS_WITH_COVERAGE:
        raise UnsupportedFormatError(f"unsupported_format: {format!r}")
    return _EXTRACTORS_WITH_COVERAGE[format](raw)
