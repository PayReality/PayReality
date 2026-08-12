"""Unit tests for text_extraction.py's Phase 3 addition
(EXPLAINABILITY_MODEL.md's Coverage Analysis): extract_text_with_coverage()
and CoverageStats. Pure, DB-free -- byte-string transformations only.
extract_text() itself is left completely unchanged and untested here
(existing coverage elsewhere is unaffected)."""

from app.domain.ai_policy_builder.text_extraction import CoverageStats, extract_text_with_coverage


def test_coverage_stats_percent_with_no_ignored_clauses():
    stats = CoverageStats(clauses_analysed=10, clauses_ignored=0, tables_extracted=0, images_skipped=0)
    assert stats.coverage_percent == 100.0


def test_coverage_stats_percent_with_some_ignored_clauses():
    stats = CoverageStats(clauses_analysed=3, clauses_ignored=1, tables_extracted=0, images_skipped=0)
    assert stats.coverage_percent == 75.0


def test_coverage_stats_percent_with_nothing_at_all_defaults_to_full():
    """Zero analysed and zero ignored (e.g. a genuinely empty document)
    is 100%, not a division-by-zero or a misleadingly low number -- there
    was nothing to miss."""
    stats = CoverageStats(clauses_analysed=0, clauses_ignored=0, tables_extracted=0, images_skipped=0)
    assert stats.coverage_percent == 100.0


def test_text_coverage_counts_one_clause_analysed():
    text, stats = extract_text_with_coverage("text", b"The CFO may approve up to $50,000.")
    assert stats.clauses_analysed == 1
    assert stats.clauses_ignored == 0
    assert stats.tables_extracted == 0
    assert stats.images_skipped == 0
    assert "$50,000" in text


def test_text_coverage_counts_empty_document_as_ignored():
    text, stats = extract_text_with_coverage("text", b"   ")
    assert stats.clauses_analysed == 0
    assert stats.clauses_ignored == 1


def test_csv_coverage_counts_rows_analysed_and_ignored():
    raw = b"role,limit\nCFO,500000\n\nManager,50000\n"
    text, stats = extract_text_with_coverage("csv", raw)
    assert stats.clauses_analysed == 3  # header + CFO row + Manager row
    assert stats.clauses_ignored == 1  # the blank line
    assert stats.tables_extracted == 1
    assert stats.images_skipped == 0


def test_xlsx_coverage_matches_text_output(tmp_path):
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["role", "limit"])
    ws.append(["CFO", 500000])
    ws.append([None, None])  # blank row, should be ignored
    path = tmp_path / "matrix.xlsx"
    wb.save(path)
    raw = path.read_bytes()

    text, stats = extract_text_with_coverage("xlsx", raw)
    assert stats.clauses_analysed == 2
    assert stats.clauses_ignored == 1
    assert stats.tables_extracted == 1
    assert "CFO" in text


def test_docx_coverage_counts_tables_and_images(tmp_path):
    from docx import Document as DocxDocument

    doc = DocxDocument()
    doc.add_paragraph("The CFO may approve up to $500,000.")
    doc.add_paragraph("")  # blank, should be ignored
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "CFO"
    table.rows[0].cells[1].text = "500000"
    path = tmp_path / "memo.docx"
    doc.save(path)
    raw = path.read_bytes()

    text, stats = extract_text_with_coverage("docx", raw)
    assert stats.clauses_analysed == 2  # 1 paragraph + 1 table row
    assert stats.clauses_ignored == 1
    assert stats.tables_extracted == 1
    assert stats.images_skipped == 0
    assert "$500,000" in text
